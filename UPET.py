# 使用前先运行
# pip install pandas sympy numpy pdfplumber openpyxl python-docx PyPDF2 pillow pywin32 openai sympy-measurement

from math import sqrt
import pandas as pd
from sympy import latex, pi, E, Symbol, sympify
import math
import re
from pathlib import Path
from openai import OpenAI
import os
import json
from sympy.parsing.latex import parse_latex
import numpy as np
import base64
import io
from PIL import ImageGrab, Image
import sys  # 新增sys模块用于系统检测
import pdfplumber
from openpyxl import Workbook
from docx.api import Document
from sympy import pi
from sympy.physics.units import *
from sympy.physics import units
from sympy_measurement import *
from sympy import symbols, diff
from urllib.parse import quote

# 新增PDF和DOCX处理库
try:
    import PyPDF2
    from docx import Document
except ImportError:
    print("缺少依赖库，请执行：pip install PyPDF2 python-docx")
    exit(1)

if sys.platform == 'win32':
    try:
        import win32com.client
        from pywintypes import com_error  # 添加COM异常处理
    except ImportError:
        print("Windows系统需要安装pywin32：pip install pywin32")
        exit(1)

CONFIG_DIR = Path.home() / ".fsearch_config"
CONFIG_FILE = CONFIG_DIR / "settings.json"
prompt_fp=f"""请对用户输入的LaTeX公式作下列修改
        对任何位置的下划线_若它的下一个字符不是左括号，则在下划线后添加成对的大括号
        在字母、数字、根号、括号外边界、分式中直接相邻的任意两者之间添加乘号，添加的乘号需要写作‘\\times ’（必须要有空格）
        除上面两点外不能做任何修改
        注意保持括号的完整性
        只能输出处理后的公式，不能输出除公式外的内容，不能输出汉字，不能输出多余字符
        公式："""
prompt_em="""提取文本所有中内容，删除文本中描述物理现象的部分，适度简化文本中叙述物理原理的部分，需要保留实验原理中涉及到数据测量的关键细节
        删除数据表格，涉及到要求的部分必须按你提取到的文字原样输出，
        思考题必须按你提取到的文字原样输出，
        细致说明实验操作方法，每一个关于测量的句子必须按你提取到的文字原样输出，保留每个数据需要测量的次数，留意涉及到计算和数据处理的部分
        输出关键计算公式，保持公式的完整性，保留每个部分最后一个公式，并根据其中含有的变量补充相应公式，使得公式构成完整的公式链条，公式中每一个物理量都必须由直接测得的物理量计算出，并展示计算出该物理量的方法，每一个出现的物理量都应该声明它是什么
        公式以LaTeX格式输出
        一次输出完整内容
        内容："""       

###############################
### 配置文件管理模块 ###
###############################
def load_config():
    """加载配置文件"""
    default_config = {
        "api_key": "",
        "default_table_path": "" 
    }
    try:
        with open(CONFIG_FILE, 'r') as f:
            loaded = json.load(f)
            return {**default_config, **loaded}
    except (FileNotFoundError, json.JSONDecodeError):
        return default_config
    
def save_config(config):
    """保存配置"""
    CONFIG_DIR.mkdir(exist_ok=True)
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=2)

def modify_default_path(config):
    """修改默认表格保存路径"""
    new_path = input("请输入新的默认表格路径：").strip('"')
    config["default_table_path"] = new_path
    save_config(config)
    print(f"默认路径已更新为：{new_path}")

def modify_API(config):
    """修改API"""
    while True:
        new_key = input("请输入新API（留空取消）：").strip()
        config["api_key"] = str(new_key)
        save_config(config)
        print("API")
        return True
    
# 修改配置菜单显示
def show_config_menu(config):
    """显示配置菜单"""
    print("\n当前配置：")
    print(f"1. DeepSeek API密钥：{config['api_key'][:4]}******")  # 安全显示API密钥
    print(f"2. 默认表格路径：{config['default_table_path'] or '未设置'}")
    print("3. 返回主菜单")
    
    choice = input("请选择要修改的项：").strip()
    if choice == '1':
        modify_API(config)
    elif choice == '2':
        modify_default_path(config)  # 新增路径修改功能
    elif choice == '3':
        return
    else:
        print("无效选项")
    

###############################
### 文件处理模块 ###
###############################
def read_excel_with_sheet_selection(file_path):
    """读取Excel文件并让用户选择工作表，返回DataFrame"""
    try:
        xls = pd.ExcelFile(file_path)
        sheets = xls.sheet_names
        
        print("\n可用工作表：")
        for i, sheet in enumerate(sheets):
            print(f"  {i+1}. {sheet}")
            
        while True:
            sheet_choice = input("请选择工作表（序号或名称，回车选择第一个）：").strip()
            if not sheet_choice:
                return pd.read_excel(xls, sheet_name=0)
            
            if sheet_choice.isdigit():
                sheet_index = int(sheet_choice) - 1
                if 0 <= sheet_index < len(sheets):
                    return pd.read_excel(xls, sheet_name=sheets[sheet_index])
                print(f"无效序号，请输入1-{len(sheets)}之间的数字")
            else:
                if sheet_choice in sheets:
                    return pd.read_excel(xls, sheet_name=sheet_choice)
                print("工作表不存在，请重新输入")
                
    except Exception as e:
        print(f"读取Excel文件失败：{str(e)}")
        raise

def extract_tables_from_file(file_path):
    """从文件中提取表格数据，返回（DataFrame列表, 错误信息）"""
    tables = []
    errors = []
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        for table in page.extract_tables():
                            if table and len(table) > 0:  # 确保有表头和至少一行数据
                                df = pd.DataFrame(table[0:])
                                tables.append(df)
                    except Exception as e:
                        errors.append(f"PDF第{page_num}页表格提取失败: {str(e)}")
        
        elif ext == '.docx':
            try:
                doc = Document(file_path)
                for table_num, table in enumerate(doc.tables, 1):
                    try:
                        data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            data.append(row_data)
                        if len(data) > 0:  # 至少需要表头和一行数据
                            df = pd.DataFrame(data[0:])
                            tables.append(df)
                    except Exception as e:
                        errors.append(f"DOCX第{table_num}个表格提取失败: {str(e)}")
            
            except Exception as e:
                errors.append(f"DOCX文件读取失败: {str(e)}")
        
        elif ext == '.doc' and sys.platform == 'win32':
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(file_path))
                for table_num, table in enumerate(doc.Tables, 1):
                    try:
                        data = []
                        for i in range(1, table.Rows.Count + 1):
                            row_data = []
                            for j in range(1, table.Columns.Count + 1):
                                cell = table.Cell(i, j)
                                text = cell.Range.Text.strip().replace('\r\x07', '')
                                row_data.append(text)
                            data.append(row_data)
                        if len(data) > 0:
                            df = pd.DataFrame(data[0:])
                            tables.append(df)
                    except Exception as e:
                        errors.append(f"DOC第{table_num}个表格提取失败: {str(e)}")
                doc.Close(SaveChanges=False)
                word.Quit()
            except Exception as e:
                errors.append(f"DOC文件处理失败: {str(e)}")
        
        else:
            errors.append(f"文件格式{ext}不支持表格提取")
    
    except Exception as e:
        errors.append(f"文件处理异常: {str(e)}")
    
    return tables, errors

def read_file_content(file_path):
    """读取不同格式文件内容"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
            return text
        elif ext == '.docx':
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif ext == '.doc':
            if sys.platform != 'win32':
                return "DOC格式仅支持Windows系统"
            
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False  # 后台运行
                doc = word.Documents.Open(os.path.abspath(file_path))
                text = doc.Content.Text
                # 清理多余的空行
                text = re.sub(r'\n\s*\n', '\n', text).strip()
                return text
            except com_error as e:
                return f"DOC文件读取失败：{e.excepinfo[2]}"
            finally:
                doc.Close(SaveChanges=False)
                word.Quit()
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return f"不支持的文件格式：{ext}"
    except Exception as e:
        return f"文件读取失败：{str(e)}"
    
def get_clickable_file_path(filename):
    # 获取基础路径
    current_file_path = os.path.abspath(__file__)
    current_dir = os.path.dirname(current_file_path)
    target_path = os.path.join(current_dir, filename)
    
    # 转换为可点击的URL格式
    file_uri = f"file://{quote(target_path, safe='/')}"  # 处理特殊字符
    # 生成终端可点击的ANSI转义序列
    return f"\033]8;;{file_uri}\033\\{target_path}\033]8;;\033\\"

    
###############################
### 单位/变量处理模块 ###
###############################
def parse_column_name(col):
    """解析列名中的变量名和单位，返回 (变量名, 单位)"""
    match = re.match(r'^([^\(]+?)\s*\((.*?)\)\s*$', col)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return col, None

def extract_var_name(full_col_name):
    """提取列名中的变量部分（去掉单位和括号）"""
    return full_col_name.split('(')[0].strip()

def extract_unit(full_col_name):
    """从列名中提取单位（返回括号内的内容，若无则返回空字符串）"""
    if '(' in full_col_name and ')' in full_col_name:
        return full_col_name.split('(')[1].split(')')[0].strip()
    return ""

def find_var_column(columns, target_var):
    for col in columns:
        # 提取括号内的单位（最多一对括号）
        unit = None
        var_part = col
        if '(' in col and ')' in col:
            # 找到第一对括号
            start = col.find('(')
            end = col.find(')')
            if start < end:
                unit = col[start+1:end]
                var_part = col[:start].strip()
            
        # 提取非中文字符作为变量名
        var_name = ''.join([c for c in var_part if ord(c) < 128 or c == '_'])
        var_name = var_name.strip()
            
        # 精确匹配
        if var_name == target_var:
            return col, unit
    return None, None

def format_unit(unit_str):
    """将单位字符串转换为LaTeX正体格式"""
    def process_part(part):
        part = re.sub(r'([a-zA-Z]+)(\^)?(\d+)?', 
                    lambda m: r'\\text{%s}%s%s' % (m.group(1), 
                    m.group(2) or '', m.group(3) or ''), part)
        return re.sub(r'\*', r'\\\\cdot ', part)
    
    if '/' in unit_str:
        numerator, denominator = unit_str.split('/', 1)
        return f"{process_part(numerator)}/{process_part(denominator)}"
    return process_part(unit_str)

def format_result_unit(unit_str):
    result_unit_latex = f"{format_unit(unit_str)}"
    replacement = f"{result_unit_latex}"
    pattern=re.escape(result_unit_latex)
    result_unit_latex = re.sub(pattern, replacement, result_unit_latex)
    return result_unit_latex

def measure_unit(unit):
    return eval(f"units.{str(unit).replace("^","**")}")


###############################
### LaTeX处理模块 ###
###############################
def simplify_latex_name(var_name):
    """将LaTeX变量名转换为可能的简化形式"""
    simplified = var_name.replace('\\', '')
    simplified = re.sub(r'_{(\w+)}', r'_\1', simplified)
    simplified = re.sub(r'\{(\w+)\}', r'\1', simplified)
    simplified = re.sub(r'\^(\w+)', r'_\1', simplified)
    return simplified

def jiagang(base_name):
    if '_' in base_name:
        prefix, suffix = base_name.split('_', 1)
        if len(prefix) >= 2 and prefix[0].isdigit()==False and prefix[0]!='-':
            modified_name = f"\\{prefix}_{suffix}"  
        else:
            modified_name =base_name
    else:
        if len(base_name) >= 2 and base_name[0].isdigit()==False and base_name[0]!='-':
            modified_name = f"\\{base_name}"  
        else: 
            modified_name =base_name
    return modified_name

def parse_latex_formula(latex_str):
    """尝试用多种方法解析LaTeX公式"""
    try:
        from latex2sympy2 import latex2sympy
        return latex2sympy(latex_str)
    except ImportError:
        return parse_latex(latex_str)

def beautify_latex_calculation(latex_str, subs_dict, var_units):
    """生成带单位格式的LaTeX算式"""

    for var, value in subs_dict.items():
        unit = var_units.get(var, None)
        
        var_name = str(var)
        # 构建正则表达式匹配完整命令
        pattern = r'(?<!\\)\\' + re.escape(var_name) + r'(?![a-zA-Z])'
        
        if unit:
            unit_tex = format_unit(unit)
            replacement = f"({value:.5g}\\ {unit_tex})"
        else:
            replacement = f"{value:.5g}"
            
        # 执行正则替换
        latex_str = re.sub(pattern, replacement, latex_str)
                    
        # 使用正则表达式替换，避免替换LaTeX命令中的字符
        # 匹配前面没有反斜杠且作为单词边界的变量名
        pattern = r'(?<!\\)\b{}'.format(re.escape(str(var)))
        latex_str = re.sub(pattern, replacement, latex_str)

    return latex_str


###############################
### 表格内容处理模块 ###
###############################
def process_cell_content(cell_content):
    """按规则处理单元格内容"""
    # 统一转换为字符串处理
    text = str(cell_content)
        
    # 步骤1：中文括号转英文括号
    text = text.replace("（", "(").replace("）", ")")
        
    # 步骤2：处理不在括号内的/
    new_text = []
    bracket_level = 0
    has_replaced = False
        
    for char in text:
        # 更新括号层级
        if char == '(': bracket_level += 1
        elif char == ')': bracket_level -= 1
            
        # 替换条件判断
        if char == '/' and bracket_level == 0:
            new_text.append('(')
            has_replaced = True
        else:
            new_text.append(char)
        
    # 拼接最终结果
    processed = ''.join(new_text)
    if has_replaced:
        processed += ')'
            
    return processed

def format_cell_value(x):
    """处理单个单元格的值"""
    if pd.isna(x):
        return " "
    elif isinstance(x, float):
        formatted = f"{x:.4f}".rstrip('0').rstrip('.')
        return formatted if formatted else "0"
    else:
        return str(x)
    
def process_content(content_str):
    """统一处理内容的方法（适用于表头和数据单元格）"""
    # 检查括号合法性
    left_count = content_str.count('(')
    right_count = content_str.count(')')
    if left_count != right_count or left_count > 1 or right_count > 1:
        raise ValueError(f"无效内容格式: {content_str}")

    # 提取单位
    unit = ''
    remaining_part = content_str
    if '(' in content_str:
        parts_before = content_str.split('(', 1)
        prefix_part = parts_before[0]
        parts_after = parts_before[1].split(')', 1)
        unit = parts_after[0]
        remaining_part = prefix_part + parts_after[1]

    # 分离中文字符和非中文字符
    chinese_part = re.sub(r'[^\u4e00-\u9fff]', '', remaining_part).strip()
    non_chinese_part = re.sub(r'[\u4e00-\u9fff]', '', remaining_part).strip()

    components = []
    if chinese_part:
        components.append(chinese_part)
    if non_chinese_part:
        var_tex = jiagang(non_chinese_part)  # 假设已定义处理函数
        components.append(f"${var_tex}$")
    if unit:
        unit_tex = format_result_unit(unit)  # 假设已定义单位处理函数
        if components:
            # 将单位添加到最后一个组件
            last_component = components[-1]
            if last_component.endswith('$'):
                components[-1] = f"{last_component[:-1]}({unit_tex})$"
            else:
                components.append(f"({unit_tex})")
        else:
            components.append(f"({unit_tex})")
    
    return ''.join(components) if components else content_str


###############################
### 核心计算模块 ###
###############################
def get_constant_value(symbol, default_value):
    """交互式获取常量的值和单位，返回 (值, 单位)"""
    constant_names = {
        pi: "π (pi)",
        E: "e (自然对数的底数)",
        Symbol('pi'): "π (pi)",
        Symbol('e'): "e (自然对数的底数)"
    }
    
    name = constant_names.get(symbol, str(symbol))
    print(f"\n检测到数学常数 {name}")
    print(f"默认值 = {default_value:.6f} ({math.floor(default_value*10000)/10000:.4f} 四舍五入后)")
    
    while True:
        choice = input("是否使用自定义值？(y/n)：").lower()
        if choice == 'y':
            try:
                custom = float(input(f"请输入 {name} 的新值："))
                unit = input(f"请输入 {name} 的单位（可选）：").strip()
                return round(custom, 4), unit if unit else None
            except:
                print("输入无效，请重新输入")
        elif choice == 'n':
            return round(default_value, 4), None
        else:
            print("请输入 y 或 n")


###############################
### AI交互模块 ###
###############################
class AISearcher:
    def __init__(self, api_key, config):
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.deepseek.com/v1"
        )
        self.file_names = set()
        self.config = config
    
    def build_prompt(self, question, mode):
        """构建带文件名限制的提示"""
        if mode == 'formula process':
            base_prompt = f"{prompt_fp}"+f"{question}"
        elif mode == 'experiment_manual':
            base_prompt = f"{prompt_em}"+f"{question}"

        return base_prompt

    def query_ai(self, question, mode):
        """执行AI查询"""
        
        response = self.client.chat.completions.create(
            model="deepseek-chat",
            messages=[{
                "role": "system",
                "content": self.build_prompt(question, mode)
            }],
            temperature=0.0
        )
        return response.choices[0].message.content.strip()


###############################
### 主功能模块 ###
###############################
def handle_table_extraction(config):
    file_path = input("\n请输入实验讲义路径（支持PDF/DOCX/TXT）（doc文件请转为docx文件再输入）\n（如果文档中表格为图片格式可能不能正确提取）: ").strip('"')
    # 自动设置默认路径逻辑
    output_dir = ""
    if config["default_table_path"]:
        use_default = input(f"是否使用默认路径 ({config['default_table_path']})? [Y/n] ").strip().lower()
        if use_default in ('', 'y', 'yes'):
            output_dir = config["default_table_path"]

    # 未使用默认路径时获取新路径
    if not output_dir:
        output_dir = input("请输入表格输出目录（留空则不保存）: ").strip('"')
        # 首次使用时自动设置默认路径
        if output_dir and not config["default_table_path"]:
            config["default_table_path"] = output_dir
            save_config(config)
            print(f"已自动保存为默认路径：{output_dir}")
    saved_path = None  # 用于存储最终保存路径

    # 表格提取和保存
    if output_dir:
        tables, errors = extract_tables_from_file(file_path)
        output_dir = Path(output_dir)
        
        if errors:
            print("\n表格提取过程中出现以下错误：")
            for error in errors:
                print(f"  - {error}")
        
        if tables:
            output_dir.mkdir(parents=True, exist_ok=True)
            base_name = Path(file_path).stem
            output_path = output_dir / f"{base_name}_tables.xlsx"
            
            try:
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    for i, df in enumerate(tables, 1):
                        # 对每个表格应用处理逻辑
                        processed_df = df.map(lambda x: process_cell_content(x))
                        processed_df.to_excel(writer, sheet_name=f"Table_{i}", index=False,header=False)
                
                saved_path = str(output_path.resolve())  # 获取绝对路径
                print(f"\n✅ 成功保存{len(tables)}个表格到：{saved_path}")
                
                # 在Windows系统自动打开资源管理器
                if sys.platform == 'win32' and os.path.exists(saved_path):
                    os.startfile(output_dir)
                
            except Exception as e:
                print(f"\n❌ 表格保存失败：{str(e)}")
        else:
            print("\n⚠️ 未找到可提取的表格")

    return file_path

specific_str = f"\\\\times "  # 这里替换为实际的特定字符串
specific_str_append=f"\\times "

def process_formula(str):
    result = process_string(str)
    result=process_string_letter(result)
    return result

def handle_spaces(text):
    # 先处理第一种情况：左右都是字母，且左侧不是特定数学函数的空格替换为指定字符串
    # 同时保留左侧是特定数学函数的空格
    # 使用正则表达式匹配左右都是字母且左侧不是特定数学函数的连续空格
    pattern1 = r'(?<!\\sin)(?<!\\cos)(?<!\\tan)(?<!\\ln)(?<!\\times)(?<=[a-zA-Z])( +)(?=[a-zA-Z])'

    # 使用正则表达式匹配左侧是特定数学函数的空格
    pattern2 = r'(\\(sin|cos|tan|ln|times)) +'

    # 处理第一种情况，替换为指定字符串
    text = re.sub(pattern1, specific_str, text)
    # 先处理左侧是特定数学函数的空格以避免干扰
    # 这里用一个临时标记替换，避免与后续替换混淆
    text = re.sub(pattern2, lambda m: m.group(1) + f"__TEMP_SPACE__{m.group(0)[-1]}", text)

    # 处理第三种情况：删除其他连续空格
    # 匹配所有剩余的连续空格
    pattern3 = r' +'
    text = re.sub(pattern3, '', text)
    # 将临时标记还原为空格
    text = re.sub(r'__TEMP_SPACE__', ' ', text)

    return text

def process_string(input_str):
    # 使用正则表达式匹配字母或数字后接反斜杠的情况，并在中间插入特定字符串
    input_str=re.sub(r'\\(sin|cos|tan|ln)',r'\\\1 ',input_str)
    input_str = handle_spaces(input_str)
    input_str=re.sub(r'_([^{])', r'_{\1}', input_str)
    input_str = re.sub(
        r'''
        (
            [A-Za-z](?<!left)  # 字母且不是left结尾
            |                  # 或
            [)}\]]              # 右圆括号或右方括号
            |                  # 或
            \d                 # 数字
        )
        ([({[])                # 左括号
        (?<!\}\{)
        (?<!\b(dfrac)\b\{)
        (?<!\b(frac)\b\{)
        (?<!\b(sqrt)\b\{)# 排除}{情况
        ''',
        r'\1' + specific_str + r'\2',
        input_str,
        flags=re.X
    )
    processed = re.sub(r'([A-Za-z0-9)}\]])\\(?!\b(right|times)\b)', r'\1' + specific_str + r'\\', input_str)
    return processed

def process_string_letter(input_str):
    # 将字符串拆分为字母块和非字母块
    tokens = re.findall(r'([A-Za-z]+|[^A-Za-z]+)', input_str)
    result = []
    prev_token = None  # 记录前一个非字母块

    for token in tokens:
        if token.isalpha():
            # 处理字母块：仅在前一个块为非字母时处理
            if prev_token is not None and not prev_token.isalpha():
                last_char = prev_token[-1]
                
                # 规则1：如果前一个非字母字符是 \，不处理
                if last_char == '\\':
                    processed = token
                else:
                    # 规则2：如果前一个非字母字符是右括号或数字
                    if last_char in {')', '}', ']'} or last_char.isdigit():
                        result.append(specific_str_append)  # 在非字母与字母块间插入
                    # 处理字母块内部：相邻字母间插入特定字符串
                    processed = specific_str_append.join(token)
            else:
                # 前一个块是字母或无前块，不处理
                processed = token
            result.append(processed)
        else:
            # 非字母块直接保留
            result.append(token)
            prev_token = token  # 更新前一个块
    
    return ''.join(result)



def main():
    config = load_config()
    full_path = get_clickable_file_path("example.xlsx")
    while True:
        query = input("\n菜单：\n1.同一公式批量数值计算；\n2.平均值；\n3.总体标准差；\n4.A类不确定度；\n5.execl表格转LaTeX；\n6.分析实验讲义；\n7.提取表格；\n8.单变量不确定度分析；\n9.合成不确定度分析；\nconfig-查看设置；\nq-退出;\n请选择：").strip().lower()
        if query == 'q':
            break

        if query == 'config':
            show_config_menu(config)
            continue

        if query == '7':
            handle_table_extraction(config)
            continue

        if query=='9':
            formula_input = input("\n请输入公式（格式如'd=\\dfrac{k\\lambda}{\\sin \\phi}',sin、cos、tan、ln前要有“\\”）\n（公式中若有连续的变量相乘，且这些变量中有变量d，请确保变量d在最右侧）：").strip()
            if '=' not in formula_input:
                print("公式格式错误，需要包含等号。")
                continue
    
            lhs, rhs = formula_input.split('=', 1)
            lhs, rhs = lhs.strip(), rhs.strip()

            rhs=process_formula(rhs)

            try:
                expr = parse_latex_formula(rhs)

            except Exception as e:
                print("解析公式时出错:", e)
                print("请检查公式格式或尝试输入SymPy格式（如'4*pi**2*l/T^2'）")
                continue

            math_constants = {
                pi: math.pi,
                E: math.e,
                Symbol('pi'): math.pi,
                Symbol('e'): math.e
            }

            constant_subs = {}
            constant_subs_cal={}
            measurements_dict={}
            var_units = {}
            variables = []
            calculation_latex=f"\\sqrt{{"
            flag=0
            for symbol in expr.free_symbols:
                if symbol in math_constants:
                    default_value = math_constants[symbol]
                    custom_value, unit = get_constant_value(symbol, default_value)
                    constant_subs[symbol] = custom_value
                    if unit:
                        var_units[symbol] = unit
                    else:
                        unit=1
                    measurements_dict[str(symbol)] = Measurement(
                    name=str(symbol),
                    values=custom_value,
                    unit=measure_unit(unit) if unit!=1 else 1,
                    uB=0
                    )
                    expr = expr.subs(symbol, measurements_dict[str(symbol)])
                else:
                    variables.append(symbol)



            file_path = input(f"""\n请输入数据文件路径（数据文件格式参考“{full_path}”表三）\n（在运行过程中不能对同一变量名重复赋值，因此修改数据后必须重启程序。同一次运行中多次出现同一变量必须保证值没有变化）：""")
            file_path=file_path.strip('"')
            try:
                if file_path.endswith(('.xlsx', '.xls')):
                    try:
                        df = read_excel_with_sheet_selection(file_path)
                    except Exception as e:
                        print("读取文件失败:", e)
                        continue
                else:
                    df = pd.read_csv(file_path, sep=None, engine='python')
            except Exception as e:
                print("读取文件失败:", e)
                continue

            data_columns = df.columns.astype(str).tolist()
            var_to_column = {}
            

            # 处理数据列匹配和自动常量检测
            for var in variables.copy():  # 使用副本遍历以便修改原列表
                original_name = str(var)
                simplified_name = simplify_latex_name(original_name)
                matches = []
        
                for col in data_columns:
                    col_var, col_unit = parse_column_name(col)
                    col_var = re.sub(r'[\u4e00-\u9fff]+', '', col_var)
                    if simplify_latex_name(col_var) == simplified_name:
                        matches.append( (col, col_unit) )

                if not matches:
                    print(f"变量 {original_name} 未找到匹配的数据列")
                    continue
                
                # 取第一个匹配项
                selected_col, selected_col_unit = matches[0]
                
                # 解析实际变量名和单位（从列头）
                name = simplified_name  # 从列头获取真实变量名
                unit_str = selected_col_unit.strip() if selected_col_unit else "1"
                
                # 转换单位符号（保持与SymPy兼容）
                try:
                    unit = eval(f"units.{unit_str.replace("^","**")}") if unit_str != "1" else 1
                except Exception as e:
                    print(f"单位解析错误，列：{selected_col}，单位：{unit_str}，错误：{str(e)}")
                    unit = 1

                # 处理列数据
                col_data = df[selected_col].dropna()  # 移除空值
                data_points = col_data.values
                
                # 提取不确定度（第二行）和数据
                uncertainty = 0.0
                data = {0.0,}
                if len(data_points) >= 1:
                    try:
                        uncertainty = float(data_points[0])  # 第一行为不确定度
                    except:
                        uncertainty = 0.0
                if len(data_points) >= 2:
                    try:
                        data = list(float(x) for x in data_points[1:])
                    except:
                        data={0}
                
                measurements_dict[name] = Measurement(
                    name=name,
                    values=data,
                    unit=unit,
                    uB=uncertainty
                )
    
                # 同时保持原有var_to_column结构
                var_to_column[var] = measurements_dict[name]
                expr = expr.subs(var, measurements_dict[name]) # 符号替换为Measurement实例
                if uncertainty!=0:
                    if flag==0:
                        flag=1
                    else:
                        calculation_latex=calculation_latex+"+"
                    calculation_latex=calculation_latex+f"\\left(\\dfrac{{\\partial {lhs}}}{{\\partial {jiagang(str(var))}}}u_{{{jiagang(str(var))}}}\\right)^2"
                
                # 从待处理变量列表中移除已匹配的变量
                variables.remove(var)

            result_unit_latex = ""

            result_unit = input("\n请输入计算结果的单位（如 kg*m/s^2，直接回车跳过）：").strip()
            if result_unit:
                result_unit_latex =format_result_unit(result_unit)
                result=convert_to(uvalue(expr)[0], measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
                u_result=convert_to(uvalue(expr)[1], measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
            else:
                result_unit_latex=""
                result=uvalue(expr)[0].evalf(5)
                u_result=uvalue(expr)[1].evalf(5)
            calculation_latex=calculation_latex+"}"

            print(f"\n结果为${result}\\ {result_unit_latex}$")
            print(f"${lhs}$的标准不确定度为：$$u_{{{lhs}}}={calculation_latex}={u_result}\\ {result_unit_latex}$$")
            print(f"相对不确定度为{uvalue(expr)[2]}（若要输出算式请选择第1项功能）\n")

        if query=='8':
            formula_input = input("\n请输入线性测量模型（格式如'l=l_0+l_1+l_2'）\n测量模型一般包括测量值，仪器误差影响，和观测产生的误差的影响：").strip()
            if '=' not in formula_input:
                print("公式格式错误，需要包含等号。")
                continue
    
            lhs, rhs = formula_input.split('=', 1)
            lhs, rhs = lhs.strip(), rhs.strip()

            try:
                expr = parse_latex_formula(rhs)

            except Exception as e:
                print("解析公式时出错:", e)
                print("请检查公式格式或尝试输入SymPy格式（如'4*pi**2*l/T^2'）")
                continue

            math_constants = {
                pi: math.pi,
                E: math.e,
                Symbol('pi'): math.pi,
                Symbol('e'): math.e
            }

            constant_subs = {}
            constant_subs_cal={}
            var_units = {}
            variables = []
            calculation_latex=f"\\sqrt{{"
            flag=0
            for symbol in expr.free_symbols:
                if symbol in math_constants:
                    default_value = math_constants[symbol]
                    custom_value, unit = get_constant_value(symbol, default_value)
                    constant_subs[symbol] = custom_value
                    if unit:
                        var_units[symbol] = unit
                else:
                    variables.append(symbol)

            expr_subs = expr.subs(constant_subs)

            file_path = input(f"""\n请输入数据文件路径（数据文件格式参考“{full_path}”表二）\n（在运行过程中不能对同一变量名重复赋值，因此修改数据后必须重启程序。同一次运行中多次出现同一变量必须保证值没有变化）：""")
            file_path=file_path.strip('"')
            try:
                if file_path.endswith(('.xlsx', '.xls')):
                    try:
                        df = read_excel_with_sheet_selection(file_path)
                    except Exception as e:
                        print("读取文件失败:", e)
                        continue
                else:
                    df = pd.read_csv(file_path, sep=None, engine='python')
            except Exception as e:
                print("读取文件失败:", e)
                continue

            data_columns = df.columns.astype(str).tolist()
            var_to_column = {}
            measurements_dict={}
            result_unit_latex = ""
            result_unit = input("\n请输入计算结果的单位（如 kg*m/s^2，直接回车跳过）：").strip()
            if result_unit:
                result_unit_latex =format_result_unit(result_unit)
            else:
                result_unit_latex=""

            # 处理数据列匹配和自动常量检测
            for var in variables.copy():  # 使用副本遍历以便修改原列表
                original_name = str(var)
                simplified_name = simplify_latex_name(original_name)
                matches = []
        
                for col in data_columns:
                    col_var, col_unit = parse_column_name(col)
                    col_var = re.sub(r'[\u4e00-\u9fff]+', '', col_var)
                    if simplify_latex_name(col_var) == simplified_name:
                        matches.append( (col, col_unit) )

                if not matches:
                    print(f"变量 {original_name} 未找到匹配的数据列")
                    continue
                
                # 取第一个匹配项
                selected_col, selected_col_unit = matches[0]
                
                # 解析实际变量名和单位（从列头）
                name = simplified_name  # 从列头获取真实变量名
                unit_str = selected_col_unit.strip() if selected_col_unit else "1"
                
                # 转换单位符号（保持与SymPy兼容）
                try:
                    unit = eval(f"units.{unit_str.replace("^","**")}") if unit_str != "1" else 1
                except Exception as e:
                    print(f"单位解析错误，列：{selected_col}，单位：{unit_str}，错误：{str(e)}")
                    unit = 1

                # 处理列数据
                col_data = df[selected_col].dropna()  # 移除空值
                data_points = col_data.values
                
                # 提取不确定度（第二行）和数据
                uncertainty = 0.0
                data = {0.0,}
                if len(data_points) >= 1:
                    try:
                        uncertainty = float(data_points[0])  # 第一行为不确定度
                    except:
                        uncertainty = 0.0 
                if len(data_points) >= 2:
                    try:
                        data = list(float(x) for x in data_points[1:])
                    except:
                        data={0}
                
                measurements_dict[name] = Measurement(
                    name=name,
                    values=data,
                    unit=unit,
                    uB=uncertainty
                )
    
                # 同时保持原有var_to_column结构
                var_to_column[var] = measurements_dict[name]
                expr = expr.subs(var, measurements_dict[name])  # 符号替换为Measurement实例
                show_name=jiagang(str(measurements_dict[name]))
                if uncertainty==0:
                    if result_unit:
                        tv=convert_to(uvalue(measurements_dict[name])[0],measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
                        uv=convert_to(uvalue(measurements_dict[name])[1],measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
                    else:
                        tv=uvalue(measurements_dict[name])[0].evalf(5)
                        uv=uvalue(measurements_dict[name])[1].evalf(5)
                    print(f"\n${show_name}$的均值为：$$\\overline{{ {show_name} }}=\\dfrac{{1}} {{{len(data)}}}\\displaystyle\\sum^{len(data)}_{{i=1}} {{{show_name}}}_i={tv}\\ {result_unit_latex}$$")
                    print(f"${show_name}$的标准不确定度为：$$u_{{ {show_name} }}=\\sqrt{{\\dfrac{{\\sum^{len(data)}_{{i=1}} ({{{show_name}}}_i-\\overline{{ {show_name} }})^2 }} {{ {len(data)} \\cdot({len(data)}-1) }} }}={uv}\\ {result_unit_latex}$$")
                else:
                    if result_unit:
                        uv=convert_to(uncertainty*unit,measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
                    else:
                        uv=uncertainty*unit.evalf(5)
                    print(f"${show_name}$的标准不确定度为：$$u_{{ {show_name} }}={uv}\\ {result_unit_latex}$$")
                if flag==0:
                    flag=1
                else:
                    calculation_latex=calculation_latex+"+"
                num=float(diff(expr,measurements_dict[name]))
                if num-int(num)==0:
                    if num==1:
                        num_str=""
                    else:
                        num_str=int(num)
                else:
                    num_str=num
                if num_str!="":
                    calculation_latex=calculation_latex+f"\\left({num_str}u_{{{jiagang(str(var))}}}\\right)^2"
                else:
                    calculation_latex=calculation_latex+f"u_{{{jiagang(str(var))}}}^2"
                
                # 从待处理变量列表中移除已匹配的变量
                variables.remove(var)
                
            if result_unit:
                result=uvalue(expr)[0]
                result=convert_to(result, measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
                u_result=convert_to(uvalue(expr)[1], measure_unit(result_unit)).as_coeff_Mul()[0].evalf(5)
            else:
                result_unit_latex=""
                result=uvalue(expr)[0].evalf(5)
                u_result=uvalue(expr)[1].evalf(5)
            calculation_latex=calculation_latex+"}"

            print(f"\n结果为${result}\\ {result_unit_latex}$")
            print(f"${lhs}$的标准不确定度为：$$u_{{{lhs}}}={calculation_latex}={u_result}\\ {result_unit_latex}$$")
            print(f"相对不确定度为{uvalue(expr)[2]}（若要输出算式请选择第1项功能）\n")



        if query == '6':
            if not config["api_key"]:
                config["api_key"] = input("请输入DeepSeek API密钥：").strip()
                save_config(config)
            
            mode=input("\n请选择处理方式:\n1.程序调用deepseek API分析讲义，并提取文档中的表格（处理时间较长且消耗较多token用量）；\n2.自行使用网页处理（推荐）；\n请选择：").strip()
            if mode=='1':
                file_path = handle_table_extraction(config)
    
                content = read_file_content(file_path)
                
                # 添加DOC文件读取结果校验
                if content.startswith("DOC文件读取失败") or content.startswith("不支持"):
                    if sys.platform == 'win32' and "权限" in content:
                        print("提示：请确保Word程序未被占用，且文件未被加密")
                    continue
                if content.startswith("文件读取失败") or content.startswith("不支持"):
                    print(content)
                    continue
                
                searcher = AISearcher(config["api_key"], config)
                try:
                    result = searcher.query_ai(content, 'experiment_manual')
                    print("\n讲义分析结果：")
                    print(result.replace(". ", ".\n").replace("- ", "\n- "))#
                except Exception as e:
                    print(f"分析失败：{str(e)}")
                continue

            elif mode=='2':
                print(f"\ndeepseek网址：https://www.deepseek.com/\n提示词：")
                print(f"       “{prompt_em}”")
                print(f"不需要开启深度思考")

            else:
                print(f"无效输入\n")
                continue

        if query == '2':
            file_path = input("\n请输入Excel文件路径: ").strip()
            file_path=file_path.strip('"')

            try:
                df = read_excel_with_sheet_selection(file_path)  # 修改后的读取方法
                df = df.drop(index=df.index[0]).reset_index(drop=True)
                # 显示读取到的列
                print("读取到列", " ".join(df.columns))
                target_vars = input("请输入变量名（用空格/逗号分隔，不需要带单位，如 rho eta_0）: ").strip()

                # 处理多种分隔符并去空
                target_list = [v.strip() for v in target_vars.replace(',', ' ').split()]
            
                # 匹配变量名并获取单位
                def process_mean(df,target_var):
                    matched_col, unit = find_var_column(df.columns.tolist(), target_var)
    
                    if not matched_col:
                        available_vars = []
                        for col in df.columns:
                            # 提取变量名部分（同上面的逻辑）
                            var_part = col
                            if '(' in col and ')' in col:
                                start = col.find('(')
                                end = col.find(')')
                                if start < end:
                                    var_part = col[:start].strip()
                            var_name = ''.join([c for c in var_part if ord(c) < 128 or c == '_'])
                            available_vars.append(var_name.strip())
                        raise ValueError(f"匹配失败！可用变量名：\n{', '.join(set(available_vars))}")
    
                    # 数据处理
                    series = pd.to_numeric(df[matched_col], errors='coerce').dropna()
        
                    if series.empty:
                        raise ValueError(f"列'{matched_col}'中没有有效数值数据")
    
                    # 计算并输出（带单位）
                    avg = round(series.mean(), 4)
                    new_var=jiagang(target_var)
                    unit_str = f" {format_result_unit(unit)}" if unit else ""
                    print(f"\n\\overline{{ {new_var} }}=\\dfrac{{1}} {{{len(series)}}}\\displaystyle\\sum^n_{{i=1}} {{{new_var}}}_i={avg}\\ {unit_str}")
                    print(f"（共使用 {len(series)} 条有效数据）")

                for var in target_list:
                    try:
                        process_mean(df, var)
                    except ValueError as e:
                        print(f"\n处理变量 {var} 时出错：{str(e)}")
                        print("----------------------------------")
                


            except Exception as e:
                print(f"\n错误: {str(e)}")

            continue

        if query == '3':
            file_path = input("\n请输入Excel文件路径: ").strip()
            file_path=file_path.strip('"')

            try:
                df = read_excel_with_sheet_selection(file_path)  # 修改后的读取方法
                df = df.drop(index=df.index[0]).reset_index(drop=True)
                # 显示读取到的列
                print("读取到列", " ".join(df.columns))
                target_vars = input("请输入变量名（用空格/逗号分隔，不需要带单位，如 rho eta_0）: ").strip()

                # 处理多种分隔符并去空
                target_list = [v.strip() for v in target_vars.replace(',', ' ').split()]
        
                def process_var(df,target_var):
                    # 匹配变量名并获取单位
                    matched_col, unit = find_var_column(df.columns.tolist(), target_var)
    
                    if not matched_col:
                        available_vars = []
                        for col in df.columns:
                        # 提取变量名部分（同上面的逻辑）
                            var_part = col
                            if '(' in col and ')' in col:
                                start = col.find('(')
                                end = col.find(')')
                                if start < end:
                                    var_part = col[:start].strip()
                            var_name = ''.join([c for c in var_part if ord(c) < 128 or c == '_'])
                            available_vars.append(var_name.strip())
                        raise ValueError(f"匹配失败！可用变量名：\n{', '.join(set(available_vars))}")
    
                    # 数据处理
                    series = pd.to_numeric(df[matched_col], errors='coerce').dropna()
            
                    if series.empty:
                        raise ValueError(f"列'{matched_col}'中没有有效数值数据")

                    # 计算并输出（带单位）
                    avg = sqrt(np.var(series))
                    new_var=jiagang(target_var)
                    unit_str = f" {format_result_unit(unit)}" if unit else ""
                    print(f"\n{{ {new_var} }}的标准差=\\sqrt{{\\dfrac{{1}} {{ {len(series)}  }} }}\\displaystyle\\sum^n_{{i=1}} ({{{new_var}}}_i-\\overline{{ {new_var} }})^2 ={avg:.2e}\\ {unit_str}")
                    print(f"（共使用 {len(series)} 条有效数据）")

                for var in target_list:
                    try:
                        process_var(df, var)
                    except ValueError as e:
                        print(f"\n处理变量 {var} 时出错：{str(e)}")
                        print("----------------------------------")

            except Exception as e:
                print(f"\n错误: {str(e)}")

            continue

        if query == '4':
            file_path = input("\n请输入Excel文件路径: ").strip()
            file_path=file_path.strip('"')

            try:
                df = read_excel_with_sheet_selection(file_path)  # 修改后的读取方法
                df = df.drop(index=df.index[0]).reset_index(drop=True)
                # 显示读取到的列
                print("读取到列", " ".join(df.columns))
                target_vars = input("请输入变量名（用空格/逗号分隔，不需要带单位，如 rho eta_0）: ").strip()

                # 处理多种分隔符并去空
                target_list = [v.strip() for v in target_vars.replace(',', ' ').split()]

                def process_A_uvalue(df,target_var):
                    # 匹配变量名并获取单位
                    matched_col, unit = find_var_column(df.columns.tolist(), target_var)
        
                    if not matched_col:
                        available_vars = []
                        for col in df.columns:
                            # 提取变量名部分（同上面的逻辑）
                            var_part = col
                            if '(' in col and ')' in col:
                                start = col.find('(')
                                end = col.find(')')
                                if start < end:
                                    var_part = col[:start].strip()
                            var_name = ''.join([c for c in var_part if ord(c) < 128 or c == '_'])
                            available_vars.append(var_name.strip())
                        raise ValueError(f"匹配失败！可用变量名：\n{', '.join(set(available_vars))}")
    
                    # 数据处理
                    series = pd.to_numeric(df[matched_col], errors='coerce').dropna()
            
                    if series.empty:
                        raise ValueError(f"列'{matched_col}'中没有有效数值数据")
    
                    # 计算并输出（带单位）
                    avg = sqrt(np.var(series)/(len(series)-1))
                    new_var=jiagang(target_var)
                    unit_str = f" {format_result_unit(unit)}" if unit else ""
                    print(f"\nu_{{ {new_var} }}=\\sqrt{{\\dfrac{{1}} {{ {len(series)} \\cdot({len(series)}-1) }} }}\\displaystyle\\sum^n_{{i=1}} ({{{new_var}}}_i-\\overline{{ {new_var} }})^2 ={avg:.2e}\\ {unit_str}")
                    print(f"（共使用 {len(series)} 条有效数据）")

                for var in target_list:
                    try:
                        process_A_uvalue(df, var)
                    except ValueError as e:
                        print(f"\n处理变量 {var} 时出错：{str(e)}")
                        print("----------------------------------")

            except Exception as e:
                print(f"\n错误: {str(e)}")

            continue

        if query =='5':
            file_path = input("\n输入数据文件路径：")
            file_path=file_path.strip('"')
            df = read_excel_with_sheet_selection(file_path)  # 修改后的读取方法

            # 显示读取到的列
            numbered_columns = [f"{i}. {col}" for i, col in enumerate(df.columns, 1)]  # 生成带序号列表
            print("  " + "\n  ".join([" ".join(numbered_columns[i:i+5]) for i in range(0, len(numbered_columns), 5)])) 
    
            selected_indices = input("输入要制表的列序号（用空格分隔，或输入'all'选择所有列）：").split()
            selected_columns = []

            # 处理全选逻辑
            if len(selected_indices) == 1 and selected_indices[0].strip().lower() in ('a', 'all'):
                selected_columns = df.columns.tolist()
            else:
                for index_str in selected_indices:
                    try:
                        index = int(index_str)
                        if index < 1 or index > len(df.columns):
                            raise ValueError
                    except ValueError:
                        raise ValueError(f"无效序号: {index_str}（有效范围：1-{len(df.columns)}）")
                    selected_columns.append(df.columns[index-1])
            
            # 用户选择方向
            while True:
                direction = input("选择表格方向\n1.原方向\n2.转置\n请选择（输入选项对应的数字序号）：").strip().lower()
                if direction in ['1', '2']:
                    break
                print("输入错误，请重新输入！")

            while True:
                line_number=input("是否添加序号\n1.否\n2.每列上方添加序号\n3.每行左侧添加序号\n请选择（输入选项对应序号）：").strip().lower()
                if line_number in ['1','2','3']:
                    break
                print("输入错误，请重新输入！")
            
            # 处理所有表头内容
            latex_vars = [process_content(str(col)).replace('Unnamed','') for col in selected_columns]
            
            # 处理有效数据行
            data = df[selected_columns]
            valid_indices = [idx for idx, row in data.iterrows() if not row.isna().all()]
            
            # 根据方向生成表格
            latex_code = []
            if direction == '1':
                # 横向表格结构
                if line_number=='1':
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * (len(selected_columns)-1)}}}",
                        "    \\hline",
                        "  " + " & ".join(latex_vars) + " \\\\",
                        "    \\hline"
                    ]
                elif line_number=='2' :
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * (len(selected_columns)-1)}}}",
                        "    \\hline",
                        " & ".join(map(str, range(1, len(selected_columns)+1))) + " \\\\",
                        "    \\hline",
                        "  " + " & ".join(latex_vars) + " \\\\",
                        "    \\hline"
                    ]
                elif line_number=='3':
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * len(selected_columns)}}}",
                        "    \\hline",
                        "     & " + " & ".join(latex_vars) + " \\\\",
                        "    \\hline"
                    ]
                
                # 填充数据行（统一处理内容）
                for row_num, idx in enumerate(valid_indices, 1):
                    values = []
                    for x in data.loc[idx]:
                        try:
                            cell_str = format_cell_value(x).replace('Unnamed','')
                            processed=process_cell_content(cell_str)
                            processed = process_content(processed)
                        except Exception as e:
                            processed = f"处理错误: {str(x)}"
                        values.append(processed)
                    if line_number=='3':
                        latex_code.append(f"    {row_num} & {' & '.join(values)} \\\\")
                    else:
                        latex_code.append(f"     {' & '.join(values)} \\\\")
                    latex_code.append("    \\hline")
                latex_code.append("\\end{tabular}")
            
            else:  # 纵向模式
                # 纵向表格结构
                if line_number=='2':
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * len(valid_indices)}}}",
                        "    \\hline",
                        "    & " + " & ".join(map(str, range(1, len(valid_indices)+1))) + " \\\\",
                        "    \\hline"
                    ]
                elif line_number=='1' :
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * len(valid_indices)}}}",
                        "    \\hline"
                    ]
                elif line_number=='3' :
                    latex_code = [
                        f"\\begin{{tabular}}{{|c|{'c|' * (len(valid_indices)+1)}}}",
                        "    \\hline"
                    ]
                
                # 填充数据行（统一处理内容）
                row_num=0
                for var, col in zip(latex_vars, selected_columns):
                    values = []
                    for idx in valid_indices:
                        x = data[col].iloc[idx]
                        try:
                            cell_str = format_cell_value(x)
                            processed=process_cell_content(cell_str)
                            processed = process_content(processed)
                        except Exception as e:
                            processed = f"处理错误: {str(x)}"
                        values.append(processed)
                    if line_number=='3':
                        row_num+=1
                        latex_code.append(f"    {row_num} & {var} & {' & '.join(values)} \\\\")
                    else:
                        latex_code.append(f"    {var} & {' & '.join(values)} \\\\")
                    latex_code.append("    \\hline")
                latex_code.append("\\end{tabular}")
            
            print("\n表格代码如下：")
            print('\n'.join(latex_code))

        if query == '1':
            formula_input = input("\n请输入公式（格式如'd=\\dfrac{k\\lambda}{\\sin \\phi}',sin、cos、tan、ln前要有“\\”）\n（公式中若有连续的变量相乘，且这些变量中有变量d，请确保变量d在最右侧）：").strip()
            if '=' not in formula_input:
                print("公式格式错误，需要包含等号。")
                continue
    
            lhs, rhs = formula_input.split('=', 1)
            lhs, rhs = lhs.strip(), rhs.strip()

            rhs=process_formula(rhs)

            try:
                expr = parse_latex_formula(rhs)

            except Exception as e:
                print("解析公式时出错:", e)
                print("请检查公式格式或尝试输入SymPy格式（如'4*pi**2*l/T^2'）")
                continue

            math_constants = {
                pi: math.pi,
                E: math.e,
                Symbol('pi'): math.pi,
                Symbol('e'): math.e
            }

            constant_subs = {}
            constant_subs_cal={}
            var_units = {}
            variables = []
            for symbol in expr.free_symbols:
                if symbol in math_constants:
                    default_value = math_constants[symbol]
                    custom_value, unit = get_constant_value(symbol, default_value)
                    constant_subs[symbol] = custom_value
                    if unit:
                        var_units[symbol] = unit
                else:
                    variables.append(symbol)

            expr_subs = expr.subs(constant_subs)

            file_path = input(f"""\n请输入数据文件路径（数据文件格式参考“{full_path}”表一）：""")
            file_path=file_path.strip('"')
            try:
                if file_path.endswith(('.xlsx', '.xls')):
                    try:
                        df = read_excel_with_sheet_selection(file_path)
                    except Exception as e:
                        print("读取文件失败:", e)
                        continue
                else:
                    df = pd.read_csv(file_path, sep=None, engine='python')
                df = df.drop(index=df.index[0]).reset_index(drop=True)
            except Exception as e:
                print("读取文件失败:", e)
                continue

            data_columns = df.columns.astype(str).tolist()
            var_to_column = {}

            # 处理数据列匹配和自动常量检测
            for var in variables.copy():  # 使用副本遍历以便修改原列表
                original_name = str(var)
                simplified_name = simplify_latex_name(original_name)
                matches = []
        
                for col in data_columns:
                    col_var, col_unit = parse_column_name(col)
                    col_var = re.sub(r'[\u4e00-\u9fff]+', '', col_var)
                    if simplify_latex_name(col_var) == simplified_name:
                        matches.append( (col, col_unit) )
        
                if len(matches) == 1:
                    selected_col, selected_unit = matches[0]
                    # 检测列数据是否唯一
                    unique_values = df[selected_col].dropna().unique()
                    if len(unique_values) == 1:
                        value = round(float(unique_values[0]), 8)
                        print(f"检测到变量 {original_name} 对应的列 '{selected_col}' 所有值为 {value}，自动视为常量")
                        constant_subs[var] = value
                        variables.remove(var)
                        if selected_unit:
                            var_units[var] = selected_unit
                            value_cal=value* measure_unit(selected_unit)
                            constant_subs_cal[var]=value_cal
                        else:
                            var_units[var]=0
                            constant_subs_cal[var]=value
                        
                        
                        continue
                
                    var_to_column[var] = selected_col
                    if selected_unit:
                        var_units[var] = selected_unit
                    else:
                        var_units[var]=0
                
                elif len(matches) > 1:
                    print(f"\n发现多个可能匹配的列：")
                    for i, (col, unit) in enumerate(matches, 1):
                        print(f"{i}. {col} (单位: {unit if unit else '无'})")
                    while True:
                        choice = input(f"请选择用于变量 {original_name} 的列（1-{len(matches)}）: ")
                        if choice.isdigit() and 1 <= int(choice) <= len(matches):
                            selected_col, selected_unit = matches[int(choice)-1]
                            # 检测列数据是否唯一
                            unique_values = df[selected_col].dropna().unique()
                            if len(unique_values) == 1:
                                value = round(float(unique_values[0]), 8)
                                print(f"检测到变量 {original_name} 对应的列 '{selected_col}' 所有值为 {value}，自动视为常量")
                                constant_subs[var] = value
                                variables.remove(var)
                                if selected_unit:
                                   var_units[var] = selected_unit
                                else:
                                    var_units[var]=0
                                expr_subs = expr_subs.subs({var: value})
                                break
                        
                            var_to_column[var] = selected_col
                            if selected_unit:
                                var_units[var] = selected_unit
                            else:
                                var_units[var]=0
                            break
                        print("输入无效，请重新输入")
                else:
                    var_name = original_name.replace('\\', '')
                    if var_name not in [parse_column_name(col)[0] for col in data_columns]:
                        while True:
                            choice = input(f"\n变量 '{var_name}' 不在数据文件中，是否为常量？(y/n): ").lower()
                            if choice == 'y':
                                try:
                                    value = float(input(f"请输入 {var_name} 的值: "))
                                    unit = input(f"请输入 {var_name} 的单位（可选）: ").strip()
                                    constant_subs[var] = round(value, 4)
                                    if unit:
                                        var_units[var] = unit
                                        constant_subs_cal[var]=constant_subs[var]*measure_unit(unit)
                                    else:
                                        constant_subs_cal[var]=constant_subs[var]
                                    variables.remove(var)
                                    
                                    break
                                except:
                                    print("输入无效，请重新输入。")
                            elif choice == 'n':
                                print(f"错误：缺少必要变量 {var_name}")
                                break
                            else:
                                print("请输入 y 或 n")

            result_unit_latex = ""
            if var_units:
                result_unit = input("\n请输入计算结果的单位（如 kg*m/s^2，直接回车跳过）：").strip()
                if result_unit:
                    result_unit_latex =format_result_unit(result_unit)
                else:
                    result_unit_latex=1

            results = []  # 新增：用于存储所有行的计算结果

            for idx, row in df.iterrows():
                subs_dict = {}
                subs_dict_cal = {}
                missing_vars = []
                for var in variables:
                    col_name = var_to_column.get(var, str(var).replace('\\', ''))
                    try:
                        subs_dict[var] = round(float(row[col_name]), 4)
                        if var_units[var]:
                            subs_dict_cal[var]=subs_dict[var]*eval(f"units.{var_units[var].replace("^","**")}")
                        else:
                            subs_dict_cal[var]=subs_dict[var]
                    except:
                        missing_vars.append(col_name)
        
                if missing_vars:
                    print(f"第{idx+1}行缺少变量：{', '.join(missing_vars)}")
                    continue

                all_subs = {**constant_subs, **subs_dict}

                try:
                    expr_subs = expr_subs.subs(constant_subs_cal)
                    expr_eval = expr_subs.subs(subs_dict_cal)
                    if result_unit:
                        result = convert_to(expr_eval, eval(f"units.{result_unit.replace("^","**")}")).as_coeff_Mul()[0].evalf(5)
                    else:result=expr_eval
                except Exception as e:
                    print(f"第{idx+1}行计算错误:", e)
                    continue

                try:
                    calculation_latex = beautify_latex_calculation(rhs, all_subs, var_units)
                    results.append({
                    'calculation_latex': calculation_latex,
                    'result': result,
                    'result_unit_latex': result_unit_latex
                })
                except Exception as e:
                    print(f"\n生成算式失败:", e)
                    print(f"{lhs} = 计算结果 = {result}\\ {result_unit_latex}")

            

            # 输出所有原始结果
            print("\n计算结果：")
            for i, res in enumerate(results):
                print(f"\n第{i+1}行: {lhs} = {res['calculation_latex']} = {res['result']}\\ {res['result_unit_latex']}")

if __name__ == "__main__":
    # 添加Windows系统检查
    if sys.platform == 'win32':
        try:
            win32com.client.Dispatch("Word.Application").Quit()
        except:
            pass
    try:
        import pandas
        from sympy import symbols
        import pdfplumber  # 新增
        from openpyxl import Workbook  # 新增
    except ImportError as e:
        print("缺少依赖库，请执行以下命令安装：")
        print("pip install pandas sympy latex2sympy2 antlr4-python3-runtime==4.11 pdfplumber openpyxl python-docx PyPDF2")
        if sys.platform == 'win32':
            print("pip install pywin32")
        exit(1)
    
    main()